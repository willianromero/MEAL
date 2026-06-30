import docx
import sys

def convert_docx_to_md(docx_path, md_path):
    doc = docx.Document(docx_path)
    with open(md_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Simple conversion based on paragraph style or text formatting
            if para.style.name.startswith('Heading 1'):
                f.write(f"\n# {text}\n")
            elif para.style.name.startswith('Heading 2'):
                f.write(f"\n## {text}\n")
            elif para.style.name.startswith('Heading 3'):
                f.write(f"\n### {text}\n")
            else:
                # Check for list items or bullet points
                if para.style.name.startswith('List') or text.startswith('•') or text.startswith('-') or (text and text[0].isdigit() and '.' in text[:3]):
                    f.write(f"\n* {text}")
                else:
                    f.write(f"\n{text}\n")
        
        # Also process tables
        for i, table in enumerate(doc.tables):
            f.write(f"\n\n### Tabla {i+1}\n\n")
            for row in table.rows:
                row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                # Write row
                f.write("| " + " | ".join(row_text) + " |\n")
                # Write header separator if first row
                if row == table.rows[0]:
                    f.write("| " + " | ".join(['---'] * len(row_text)) + " |\n")

if __name__ == '__main__':
    convert_docx_to_md('ANEXO TÉCNICO CLINICA MAICAO.docx', 'ANEXO_TECNICO_CLINICA_MAICAO.md')
    print("Conversión completada con éxito.")
